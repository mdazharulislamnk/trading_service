import os
import sys

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Final Exam Report - CF/PNTL 06 Hafiz', 0)

# Question 3
doc.add_heading('Question-03: MS Excel Task', level=1)
doc.add_paragraph('Calculated Salary Sheet based on 30% House Rent, 10% Medical, 8% Transport:')

table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Employee ID'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Basic Pay'
hdr_cells[3].text = 'House Rent'
hdr_cells[4].text = 'Medical Allowance'
hdr_cells[5].text = 'Transport'
hdr_cells[6].text = 'Gross Pay'

records = [
    ('E-01', 'Alice Smith', '20000', '6000', '2000', '1600', '29600'),
    ('E-02', 'Bob Johnson', '25000', '7500', '2500', '2000', '37000'),
    ('E-03', 'Charlie Brown', '15000', '4500', '1500', '1200', '22200'),
    ('E-04', 'Diana Prince', '30000', '9000', '3000', '2400', '44400'),
    ('E-05', 'Edward Kenway', '18000', '5400', '1800', '1440', '26640')
]

for rec in records:
    row_cells = table.add_row().cells
    for i, val in enumerate(rec):
        row_cells[i].text = val

doc.add_paragraph('\n')

# Question 4
doc.add_heading('Question-04: MS Access Task', level=1)
doc.add_paragraph('Table: Selection')

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Trainee ID'
hdr2[1].text = 'Name'
hdr2[2].text = 'Chose GAVE'
hdr2[3].text = 'Has Experience'

records2 = [
    ('1257138', 'Afsana Akter', 'Yes', 'No'),
    ('1258749', 'Md Tahsin Hasan', 'Yes', 'Yes'),
    ('1258929', 'TASLIMA AKTER', 'No', 'No'),
    ('1258941', 'Md.Sajib', 'Yes', 'Yes'),
    ('1258942', 'Md Mojib Ullah', 'No', 'No'),
    ('1259202', 'MD.BIPLAB HOSSAIN', 'No', 'Yes')
]

for rec in records2:
    row2 = table2.add_row().cells
    for i, val in enumerate(rec):
        row2[i].text = val

doc.add_paragraph('\nQuery Result (Chose GAVE = Yes & Has Experience = Yes):')

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Trainee ID'
hdr3[1].text = 'Name'
hdr3[2].text = 'Chose GAVE'
hdr3[3].text = 'Has Experience'

records3 = [
    ('1258749', 'Md Tahsin Hasan', 'Yes', 'Yes'),
    ('1258941', 'Md.Sajib', 'Yes', 'Yes')
]

for rec in records3:
    row3 = table3.add_row().cells
    for i, val in enumerate(rec):
        row3[i].text = val

output_path = r'C:\Users\AZHAR\Desktop\CF-PNTL 06 Hafiz\Report_Q3_Q4.docx'

# Create dir if not exist
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print("MS Word Report created successfully at: " + output_path)
